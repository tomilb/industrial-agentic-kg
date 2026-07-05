import re
from pathlib import Path

import pytest
import server
from docx import Document
from docx.oxml.ns import qn

# --- Driver / sesión falsos (sin Neo4j real) --------------------------------


class _FakeResult:
    def __init__(self, single_valor=None, data_valor=None):
        self._single_valor = single_valor
        self._data_valor = data_valor if data_valor is not None else []

    def single(self):
        return self._single_valor

    def data(self):
        return self._data_valor


class _FakeSession:
    def __init__(self, respuestas_compartidas):
        self._respuestas = respuestas_compartidas  # cola compartida, no se copia

    def run(self, query, **kwargs):
        return self._respuestas.pop(0)

    def __enter__(self):
        return self

    def __exit__(self, *exc_info):
        return False


class _FakeDriver:
    """El driver posee la cola mutable de respuestas; cada sesión que abre
    comparte la MISMA lista (por referencia), así que varias llamadas a
    driver.session() seguidas consumen la cola en orden, igual que pasaría
    con un driver real y varias queries secuenciales."""

    def __init__(self, respuestas):
        self._respuestas = list(respuestas)

    def session(self):
        return _FakeSession(self._respuestas)


class _FakeSessionQueRompe(_FakeSession):
    def run(self, query, **kwargs):
        raise RuntimeError("Neo4j no disponible")


class _FakeDriverQueRompe:
    def session(self):
        return _FakeSessionQueRompe([])


MAQUINA_M3 = {
    "id": "M3",
    "nombre": "AOI de gama básica",
    "capacidad_ud_h": 400.0,
    "coste_adquisicion": 70_000.0,
    "coste_mantenimiento_anual": 5_000.0,
    "consumo_kwh": 3.0,
    "vida_util_anos": 8,
    "oee_actual": 0.75,
}

CANDIDATA_C3 = {
    "id": "C3",
    "nombre": "AOI con IA de nueva generación",
    "capacidad_ud_h": 650.0,
    "coste_adquisicion": 150_000.0,
    "coste_mantenimiento_anual": 5_000.0,
    "consumo_kwh": 3.0,
    "vida_util_anos": 8,
    "mejora_capacidad_pct": 62.0,
}


# --- _ordenar_estaciones_por_cadena -----------------------------------------


def test_ordenar_estaciones_por_cadena_reconstruye_el_orden():
    filas_desordenadas = [
        {"id": "E3", "nombre": "AOI", "maquina_id": "M3", "predecesora_id": "E2"},
        {"id": "E1", "nombre": "SMD", "maquina_id": "M1", "predecesora_id": None},
        {"id": "E5", "nombre": "Test", "maquina_id": "M5", "predecesora_id": "E4"},
        {"id": "E2", "nombre": "Reflow", "maquina_id": "M2", "predecesora_id": "E1"},
        {"id": "E4", "nombre": "Carcasa", "maquina_id": "M4", "predecesora_id": "E3"},
    ]

    ordenadas = server._ordenar_estaciones_por_cadena(filas_desordenadas)

    assert [e["id"] for e in ordenadas] == ["E1", "E2", "E3", "E4", "E5"]


# --- _consultar_specs_maquina ------------------------------------------------


def test_specs_maquina_ok():
    driver = _FakeDriver([_FakeResult(single_valor={"maquina": MAQUINA_M3})])

    resultado = server._consultar_specs_maquina(driver, {"maquina_id": "M3"})

    assert resultado.maquina.id == "M3"
    assert resultado.maquina.oee_actual == 0.75
    assert resultado.maquina.mejora_capacidad_pct is None


def test_specs_maquina_no_encontrada():
    driver = _FakeDriver([_FakeResult(single_valor=None)])

    with pytest.raises(ValueError, match="No existe"):
        server._consultar_specs_maquina(driver, {"maquina_id": "M99"})


def test_specs_maquina_sin_parametro():
    with pytest.raises(ValueError, match="Falta el parámetro"):
        server._consultar_specs_maquina(_FakeDriver([]), {})


# --- _consultar_topologia_linea ----------------------------------------------


def test_topologia_linea_ok():
    filas = [
        {"id": "E1", "nombre": "SMD", "maquina_id": "M1", "predecesora_id": None},
        {"id": "E2", "nombre": "Reflow", "maquina_id": "M2", "predecesora_id": "E1"},
    ]
    driver = _FakeDriver(
        [
            _FakeResult(single_valor={"id": "L1"}),
            _FakeResult(data_valor=filas),
        ]
    )

    resultado = server._consultar_topologia_linea(driver, {})

    assert resultado.linea_id == "L1"
    assert [e.id for e in resultado.estaciones] == ["E1", "E2"]


def test_topologia_linea_sin_lineas_en_el_grafo():
    driver = _FakeDriver([_FakeResult(single_valor=None)])

    with pytest.raises(ValueError, match="No hay ninguna Linea"):
        server._consultar_topologia_linea(driver, {})


# --- _consultar_candidatos_sustitucion ---------------------------------------


def test_candidatos_sustitucion_ok():
    driver = _FakeDriver(
        [
            _FakeResult(single_valor={"maquina": MAQUINA_M3}),
            _FakeResult(data_valor=[{"candidata": CANDIDATA_C3}]),
        ]
    )

    resultado = server._consultar_candidatos_sustitucion(driver, {"maquina_id": "M3"})

    assert resultado.maquina_actual.id == "M3"
    assert len(resultado.candidatas) == 1
    assert resultado.candidatas[0].mejora_capacidad_pct == 62.0


def test_candidatos_sustitucion_sin_candidatas():
    driver = _FakeDriver(
        [
            _FakeResult(single_valor={"maquina": MAQUINA_M3}),
            _FakeResult(data_valor=[]),
        ]
    )

    resultado = server._consultar_candidatos_sustitucion(driver, {"maquina_id": "M3"})

    assert resultado.candidatas == []


def test_candidatos_sustitucion_maquina_no_encontrada():
    driver = _FakeDriver([_FakeResult(single_valor=None)])

    with pytest.raises(ValueError, match="No existe"):
        server._consultar_candidatos_sustitucion(driver, {"maquina_id": "M99"})


# --- _ejecutar_consulta (dispatcher + contrato de error) --------------------


def test_ejecutar_consulta_pregunta_no_reconocida_devuelve_error_estructurado():
    resultado = server._ejecutar_consulta(_FakeDriver([]), "pregunta_inventada", {})

    assert resultado["error"] == "pregunta_estructurada no reconocida"
    assert "pregunta_inventada" in resultado["detalle"]


def test_ejecutar_consulta_maquina_no_encontrada_devuelve_error_estructurado():
    driver = _FakeDriver([_FakeResult(single_valor=None)])

    resultado = server._ejecutar_consulta(driver, "specs_maquina", {"maquina_id": "M99"})

    assert "No existe" in resultado["error"]
    assert resultado["detalle"] is None


def test_ejecutar_consulta_excepcion_neo4j_devuelve_error_estructurado():
    resultado = server._ejecutar_consulta(
        _FakeDriverQueRompe(), "specs_maquina", {"maquina_id": "M3"}
    )

    assert resultado["error"] == "Error al consultar el grafo"
    assert "Neo4j no disponible" in resultado["detalle"]


def test_ejecutar_consulta_ok_devuelve_dict_serializable():
    driver = _FakeDriver([_FakeResult(single_valor={"maquina": MAQUINA_M3})])

    resultado = server._ejecutar_consulta(driver, "specs_maquina", {"maquina_id": "M3"})

    assert isinstance(resultado, dict)
    assert resultado["maquina"]["id"] == "M3"
    assert resultado["maquina"]["oee_actual"] == 0.75


# --- calcular_financiero -----------------------------------------------------


def test_calcular_financiero_caso_feliz():
    entrada = server.CalcularFinancieroInput(
        coste_inversion=300_000,
        capacidad_actual_ud_h=100,
        capacidad_nueva_ud_h=150,
        margen_por_unidad=10,
        horas_operacion_dia=10,
        horizonte_anos=3,
    )

    resultado = server._calcular_financiero(entrada)

    ganancia_diaria = 50 * 10 * 10  # 5000
    payback_esperado = 300_000 / (ganancia_diaria * 30)
    roi_esperado = ((ganancia_diaria * 365 * 3) - 300_000) / 300_000 * 100
    van_esperado = -300_000 + sum((ganancia_diaria * 365) / (1.10**anio) for anio in range(1, 4))

    assert resultado.payback_meses == pytest.approx(payback_esperado)
    assert resultado.roi_pct == pytest.approx(roi_esperado)
    assert resultado.van == pytest.approx(van_esperado)


def test_capacidad_nueva_igual_a_la_actual_no_es_error():
    entrada = server.CalcularFinancieroInput(
        coste_inversion=100_000,
        capacidad_actual_ud_h=100,
        capacidad_nueva_ud_h=100,
        margen_por_unidad=10,
        horas_operacion_dia=10,
        horizonte_anos=5,
    )

    resultado = server._calcular_financiero(entrada)

    assert resultado.payback_meses is None
    assert resultado.roi_pct == pytest.approx(-100.0)
    assert resultado.van == pytest.approx(-100_000)


def test_capacidad_nueva_peor_que_la_actual_no_es_error():
    entrada = server.CalcularFinancieroInput(
        coste_inversion=300_000,
        capacidad_actual_ud_h=150,
        capacidad_nueva_ud_h=100,
        margen_por_unidad=10,
        horas_operacion_dia=10,
        horizonte_anos=3,
    )

    resultado = server._calcular_financiero(entrada)

    assert resultado.payback_meses is None
    assert resultado.roi_pct < -100.0
    assert resultado.van < -300_000


def test_margen_por_unidad_cero_devuelve_error_estructurado():
    resultado = server._ejecutar_calculo_financiero(
        dict(
            coste_inversion=100_000,
            capacidad_actual_ud_h=100,
            capacidad_nueva_ud_h=150,
            margen_por_unidad=0,
            horas_operacion_dia=10,
            horizonte_anos=3,
        )
    )

    assert resultado["error"] == "Parámetros de entrada inválidos"
    assert resultado["detalle"] is not None


def test_margen_por_unidad_negativo_devuelve_error_estructurado():
    resultado = server._ejecutar_calculo_financiero(
        dict(
            coste_inversion=100_000,
            capacidad_actual_ud_h=100,
            capacidad_nueva_ud_h=150,
            margen_por_unidad=-5,
            horas_operacion_dia=10,
            horizonte_anos=3,
        )
    )

    assert resultado["error"] == "Parámetros de entrada inválidos"


def test_coste_inversion_no_positivo_devuelve_error_estructurado():
    resultado = server._ejecutar_calculo_financiero(
        dict(
            coste_inversion=0,
            capacidad_actual_ud_h=100,
            capacidad_nueva_ud_h=150,
            margen_por_unidad=10,
            horas_operacion_dia=10,
            horizonte_anos=3,
        )
    )

    assert resultado["error"] == "Parámetros de entrada inválidos"


def test_payback_supera_el_horizonte_no_se_recorta():
    entrada = server.CalcularFinancieroInput(
        coste_inversion=1_000_000,
        capacidad_actual_ud_h=100,
        capacidad_nueva_ud_h=110,
        margen_por_unidad=1,
        horas_operacion_dia=8,
        horizonte_anos=1,
    )

    resultado = server._calcular_financiero(entrada)

    ganancia_diaria = 10 * 1 * 8  # 80
    payback_esperado = 1_000_000 / (ganancia_diaria * 30)

    assert resultado.payback_meses == pytest.approx(payback_esperado)
    assert resultado.payback_meses > entrada.horizonte_anos * 12


def test_ejecutar_calculo_financiero_excepcion_inesperada_devuelve_error_estructurado(
    monkeypatch,
):
    def _romper(entrada):
        raise RuntimeError("fallo inesperado")

    monkeypatch.setattr(server, "_calcular_financiero", _romper)

    resultado = server._ejecutar_calculo_financiero(
        dict(
            coste_inversion=100_000,
            capacidad_actual_ud_h=100,
            capacidad_nueva_ud_h=150,
            margen_por_unidad=10,
            horas_operacion_dia=10,
            horizonte_anos=3,
        )
    )

    assert resultado["error"] == "Error al calcular la viabilidad financiera"
    assert "fallo inesperado" in resultado["detalle"]


# --- detectar_cuello_botella --------------------------------------------------


def test_calcular_oee_y_utilizacion_con_registros():
    registros = [
        {"unidades": 4800, "paradas_min": 0.0},
        {"unidades": 3200, "paradas_min": 480.0},
    ]

    oee, utilizacion_pct = server._calcular_oee_y_utilizacion(registros, capacidad_ud_h=400.0)

    # dia 1: 4800/(400*16)=0.75, disponibilidad 100%
    # dia 2: 3200/(400*16)=0.50, disponibilidad (960-480)/960*100=50%
    assert oee == pytest.approx((0.75 + 0.50) / 2)
    assert utilizacion_pct == pytest.approx((100.0 + 50.0) / 2)


def test_calcular_oee_y_utilizacion_sin_registros_da_none():
    oee, utilizacion_pct = server._calcular_oee_y_utilizacion([], capacidad_ud_h=400.0)

    assert oee is None
    assert utilizacion_pct is None


def test_resolver_oee_para_capacidad_usa_empirico_si_existe():
    oee, fuente = server._resolver_oee_para_capacidad(0.6, oee_actual=0.85)

    assert oee == 0.6
    assert fuente == "empirico"


def test_resolver_oee_para_capacidad_usa_estatico_si_empirico_es_none():
    oee, fuente = server._resolver_oee_para_capacidad(None, oee_actual=0.85)

    assert oee == 0.85
    assert fuente == "estatico_sin_datos"


def test_detectar_cuello_botella_usa_capacidad_efectiva_no_nominal():
    # E1 tiene capacidad nominal MUY superior a E2, pero un oee empírico tan
    # bajo que su capacidad_efectiva es la menor -> debe ser la restricción.
    # Las unidades se eligen para que el oee empírico coincida exactamente
    # con oee_actual, así los números de capacidad_efectiva no cambian
    # respecto a antes, pero ahora pasan por la vía "empirico".
    rows = [
        {
            "estacion_id": "E1",
            "estacion_nombre": "Estación grande pero ineficiente",
            "maquina_id": "M1",
            "capacidad_ud_h": 1000.0,
            "oee_actual": 0.2,
            "registros": [{"unidades": 3200, "paradas_min": 0.0}],  # 3200/(1000*16)=0.2
        },
        {
            "estacion_id": "E2",
            "estacion_nombre": "Estación pequeña pero eficiente",
            "maquina_id": "M2",
            "capacidad_ud_h": 300.0,
            "oee_actual": 0.9,
            "registros": [{"unidades": 4320, "paradas_min": 0.0}],  # 4320/(300*16)=0.9
        },
    ]
    driver = _FakeDriver([_FakeResult(data_valor=rows)])

    resultado = server._detectar_cuello_botella(driver, "L1", "2026-01-01", "2026-01-31")

    por_id = {e.estacion_id: e for e in resultado.estaciones}
    assert por_id["E1"].capacidad_efectiva == pytest.approx(200.0)
    assert por_id["E2"].capacidad_efectiva == pytest.approx(270.0)
    assert por_id["E1"].fuente_oee == "empirico"
    assert por_id["E2"].fuente_oee == "empirico"
    assert por_id["E1"].es_restriccion is True
    assert por_id["E2"].es_restriccion is False


def test_detectar_cuello_botella_oee_empirico_invierte_la_restriccion_respecto_al_estatico():
    # Estación A: ficha buena (oee_actual=0.85, efectiva estática=510), pero
    # rindiendo muy por debajo en la práctica (oee empírico=0.40 -> efectiva=240).
    # AOI: ficha peor (oee_actual=0.75, efectiva estática=300), pero rindiendo
    # exactamente como su ficha (oee empírico=0.75 -> efectiva=300).
    #
    # Con el cálculo ESTÁTICO puro (solo oee_actual): AOI (300) < A (510) ->
    # AOI sería la restricción. Con el cálculo EMPÍRICO: A (240) < AOI (300)
    # -> A pasa a ser la restricción. Este test confirma que la tool usa el
    # segundo criterio.
    rows = [
        {
            "estacion_id": "EA",
            "estacion_nombre": "Estación A (ficha buena, rendimiento real malo)",
            "maquina_id": "MA",
            "capacidad_ud_h": 600.0,
            "oee_actual": 0.85,
            "registros": [{"unidades": 3840, "paradas_min": 0.0}],  # 3840/(600*16)=0.40
        },
        {
            "estacion_id": "EAOI",
            "estacion_nombre": "AOI",
            "maquina_id": "M3",
            "capacidad_ud_h": 400.0,
            "oee_actual": 0.75,
            "registros": [{"unidades": 4800, "paradas_min": 0.0}],  # 4800/(400*16)=0.75
        },
    ]
    driver = _FakeDriver([_FakeResult(data_valor=rows)])

    resultado = server._detectar_cuello_botella(driver, "L1", "2026-01-01", "2026-01-31")

    por_id = {e.estacion_id: e for e in resultado.estaciones}

    capacidad_estatica_a = rows[0]["capacidad_ud_h"] * rows[0]["oee_actual"]
    capacidad_estatica_aoi = rows[1]["capacidad_ud_h"] * rows[1]["oee_actual"]
    assert capacidad_estatica_aoi < capacidad_estatica_a  # con el criterio viejo, AOI ganaba

    assert por_id["EA"].oee == pytest.approx(0.40)
    assert por_id["EA"].capacidad_efectiva == pytest.approx(240.0)
    assert por_id["EAOI"].capacidad_efectiva == pytest.approx(300.0)
    assert por_id["EA"].es_restriccion is True
    assert por_id["EAOI"].es_restriccion is False


def test_detectar_cuello_botella_empate_marca_ambas():
    rows = [
        {
            "estacion_id": "E1",
            "estacion_nombre": "Est1",
            "maquina_id": "M1",
            "capacidad_ud_h": 400.0,
            "oee_actual": 0.75,
            "registros": [],
        },
        {
            "estacion_id": "E2",
            "estacion_nombre": "Est2",
            "maquina_id": "M2",
            "capacidad_ud_h": 500.0,
            "oee_actual": 0.6,
            "registros": [],
        },
        {
            "estacion_id": "E3",
            "estacion_nombre": "Est3",
            "maquina_id": "M3",
            "capacidad_ud_h": 600.0,
            "oee_actual": 0.9,
            "registros": [],
        },
    ]
    driver = _FakeDriver([_FakeResult(data_valor=rows)])

    resultado = server._detectar_cuello_botella(driver, "L1", "2026-01-01", "2026-01-31")

    por_id = {e.estacion_id: e for e in resultado.estaciones}
    assert por_id["E1"].es_restriccion is True
    assert por_id["E2"].es_restriccion is True
    assert por_id["E3"].es_restriccion is False
    assert all(e.fuente_oee == "estatico_sin_datos" for e in resultado.estaciones)


def test_detectar_cuello_botella_sin_registros_en_el_periodo_da_oee_none():
    rows = [
        {
            "estacion_id": "E1",
            "estacion_nombre": "Est1",
            "maquina_id": "M1",
            "capacidad_ud_h": 400.0,
            "oee_actual": 0.75,
            "registros": [],
        },
    ]
    driver = _FakeDriver([_FakeResult(data_valor=rows)])

    resultado = server._detectar_cuello_botella(driver, "L1", "1999-01-01", "1999-01-10")

    estacion = resultado.estaciones[0]
    assert estacion.capacidad_efectiva == pytest.approx(300.0)
    assert estacion.es_restriccion is True
    assert estacion.oee is None
    assert estacion.utilizacion_pct is None
    assert estacion.fuente_oee == "estatico_sin_datos"


def test_detectar_cuello_botella_periodo_mal_formado_devuelve_error_estructurado():
    resultado = server._ejecutar_deteccion_cuello_botella(
        _FakeDriver([]), "L1", "no-es-una-fecha", "2026-01-10"
    )

    assert "Fecha de periodo inválida" in resultado["error"]


def test_detectar_cuello_botella_desde_posterior_a_hasta_devuelve_error_estructurado():
    resultado = server._ejecutar_deteccion_cuello_botella(
        _FakeDriver([]), "L1", "2026-02-01", "2026-01-01"
    )

    assert "posterior" in resultado["error"]


def test_detectar_cuello_botella_linea_inexistente_devuelve_error_estructurado():
    driver = _FakeDriver([_FakeResult(data_valor=[])])

    resultado = server._ejecutar_deteccion_cuello_botella(driver, "L99", "2026-01-01", "2026-01-31")

    assert "No existe" in resultado["error"]


def test_ejecutar_deteccion_excepcion_inesperada_devuelve_error_estructurado(monkeypatch):
    def _romper(driver, linea_id, desde, hasta):
        raise RuntimeError("Neo4j no disponible")

    monkeypatch.setattr(server, "_detectar_cuello_botella", _romper)

    resultado = server._ejecutar_deteccion_cuello_botella(
        _FakeDriver([]), "L1", "2026-01-01", "2026-01-31"
    )

    assert resultado["error"] == "Error al detectar el cuello de botella"
    assert "Neo4j no disponible" in resultado["detalle"]


# --- generar_informe ----------------------------------------------------------

MAQUINA_GENERICA = {
    "id": "MX",
    "nombre": "Máquina genérica",
    "capacidad_ud_h": 600.0,
    "coste_adquisicion": 100_000.0,
    "coste_mantenimiento_anual": 5_000.0,
    "consumo_kwh": 5.0,
    "vida_util_anos": 10,
    "oee_actual": 0.85,
}


def test_calcular_rango_periodo_mensual_y_trimestral():
    ref = server.date(2026, 7, 4)

    desde, hasta = server._calcular_rango_periodo("mensual", ref)
    assert hasta == ref
    assert (hasta - desde).days + 1 == 30

    desde3, hasta3 = server._calcular_rango_periodo("trimestral", ref)
    assert hasta3 == ref
    assert (hasta3 - desde3).days + 1 == 91


def test_calcular_rango_periodo_invalido_lanza_valueerror():
    with pytest.raises(ValueError):
        server._calcular_rango_periodo("semanal", server.date(2026, 7, 4))


def test_calcular_rango_periodo_anterior_contiguo_misma_duracion():
    desde = server.date(2026, 6, 5)
    hasta = server.date(2026, 7, 4)  # 30 días

    desde_ant, hasta_ant = server._calcular_rango_periodo_anterior(desde, hasta)

    assert hasta_ant == desde - server.timedelta(days=1)
    assert (hasta_ant - desde_ant).days + 1 == (hasta - desde).days + 1


def test_serie_produccion_diaria_suma_por_fecha():
    rows = [
        {
            "registros": [
                {"fecha": "2026-01-01", "unidades": 100},
                {"fecha": "2026-01-02", "unidades": 110},
            ]
        },
        {
            "registros": [
                {"fecha": "2026-01-01", "unidades": 200},
                {"fecha": "2026-01-02", "unidades": 210},
            ]
        },
    ]

    serie = server._serie_produccion_diaria(rows)

    assert serie == [("2026-01-01", 300), ("2026-01-02", 320)]


def test_calcular_kpis_estacion_totales_correctos():
    rows = [
        {
            "estacion_id": "E1",
            "estacion_nombre": "Est1",
            "registros": [
                {"fecha": "2026-01-01", "unidades": 100, "defectos": 5, "paradas_min": 0.0},
                {"fecha": "2026-01-02", "unidades": 110, "defectos": 3, "paradas_min": 0.0},
            ],
        },
    ]
    estaciones = [{"estacion_id": "E1", "oee": 0.8, "utilizacion_pct": 95.0}]

    kpis = server._calcular_kpis_estacion(rows, estaciones)

    assert kpis == [
        {
            "estacion_nombre": "Est1",
            "utilizacion_pct": 95.0,
            "oee": 0.8,
            "unidades_totales": 210,
            "defectos_totales": 8,
        }
    ]


def test_formato_numero_es_separador_de_miles():
    assert server._formato_numero_es(2234516) == "2.234.516"
    assert server._formato_numero_es(9500) == "9.500"
    assert server._formato_numero_es(210) == "210"


def test_formato_numero_es_con_decimales_usa_coma():
    assert server._formato_numero_es(0.104, 1) == "0,1"
    assert server._formato_numero_es(34940.0) == "34.940"
    assert server._formato_numero_es(0.75, 2) == "0,75"


def test_generar_diagnostico_y_recomendaciones_con_candidata_cita_payback_real():
    restriccion = {
        "estacion_id": "E3",
        "estacion_nombre": "AOI de gama básica",
        "maquina_id": "M3",
        "capacidad_ud_h": 400.0,
        "oee_actual": 0.75,
        "capacidad_efectiva": 300.0,
        "oee": 0.75,
        "utilizacion_pct": 95.0,
        "fuente_oee": "empirico",
        "es_restriccion": True,
    }
    segunda = {
        "estacion_id": "E4",
        "estacion_nombre": "Cobot",
        "maquina_id": "M4",
        "capacidad_ud_h": 500.0,
        "oee_actual": 0.88,
        "capacidad_efectiva": 440.0,
        "oee": 0.88,
        "utilizacion_pct": 95.0,
        "fuente_oee": "empirico",
        "es_restriccion": False,
    }
    estaciones_actual = [restriccion, segunda]
    estaciones_anterior = [dict(restriccion), dict(segunda)]  # sin cambios

    driver = _FakeDriver(
        [
            _FakeResult(single_valor={"maquina": MAQUINA_M3}),
            _FakeResult(data_valor=[{"candidata": CANDIDATA_C3}]),
        ]
    )

    diagnostico, recomendaciones = server._generar_diagnostico_y_recomendaciones(
        driver, estaciones_actual, estaciones_anterior
    )

    assert "AOI de gama básica" in diagnostico
    assert len(recomendaciones) == 2
    assert "sigue siendo la restricción" in recomendaciones[0]
    assert "payback estimado" in recomendaciones[1]
    assert "ROI a 3 años" in recomendaciones[1]
    # nombre legible de la máquina, nunca el id interno
    assert "AOI de gama básica" in recomendaciones[1]
    assert "M3" not in recomendaciones[1]
    # el diagnóstico no debe repetirse literalmente dentro de las recomendaciones
    assert diagnostico not in recomendaciones


def test_generar_recomendaciones_sin_candidata_y_oee_bajo_recomienda_mantenimiento():
    restriccion = {
        "estacion_id": "EX",
        "estacion_nombre": "Estación genérica",
        "maquina_id": "MX",
        "capacidad_ud_h": 600.0,
        "oee_actual": 0.85,
        "capacidad_efectiva": 360.0,
        "oee": 0.60,
        "utilizacion_pct": 80.0,
        "fuente_oee": "empirico",
        "es_restriccion": True,
    }
    estaciones_actual = [restriccion]
    estaciones_anterior = [dict(restriccion)]

    driver = _FakeDriver(
        [
            _FakeResult(single_valor={"maquina": MAQUINA_GENERICA}),
            _FakeResult(data_valor=[]),
        ]
    )

    _, recomendaciones = server._generar_diagnostico_y_recomendaciones(
        driver, estaciones_actual, estaciones_anterior
    )

    assert any("mantenimiento" in r for r in recomendaciones)


def test_generar_recomendaciones_cambio_de_restriccion_lo_indica():
    ea_actual = {
        "estacion_id": "EA",
        "estacion_nombre": "Estación A",
        "maquina_id": "MA",
        "capacidad_ud_h": 600.0,
        "oee_actual": 0.85,
        "capacidad_efectiva": 240.0,
        "oee": 0.40,
        "utilizacion_pct": 80.0,
        "fuente_oee": "empirico",
        "es_restriccion": True,
    }
    eb_actual = {
        "estacion_id": "EB",
        "estacion_nombre": "Estación B",
        "maquina_id": "MB",
        "capacidad_ud_h": 400.0,
        "oee_actual": 0.75,
        "capacidad_efectiva": 300.0,
        "oee": 0.75,
        "utilizacion_pct": 95.0,
        "fuente_oee": "empirico",
        "es_restriccion": False,
    }
    ea_anterior = {**ea_actual, "capacidad_efectiva": 500.0, "es_restriccion": False}
    eb_anterior = {**eb_actual, "capacidad_efectiva": 200.0, "es_restriccion": True}

    driver = _FakeDriver(
        [
            _FakeResult(single_valor={"maquina": MAQUINA_GENERICA}),  # candidatos_sustitucion de MA
            _FakeResult(data_valor=[]),
        ]
    )

    _, recomendaciones = server._generar_diagnostico_y_recomendaciones(
        driver, [ea_actual, eb_actual], [ea_anterior, eb_anterior]
    )

    assert "ha cambiado" in recomendaciones[0]
    assert "Estación B" in recomendaciones[0]
    assert "Estación A" in recomendaciones[0]


def _construir_datos_informe_fixture() -> tuple[list[dict], list[dict]]:
    """rows_actual/rows_anterior para 2 estaciones (E3/M3=AOI, restricción
    con candidata C3; E4/M4=Ensamblaje), reutilizados por los tests de
    generar_informe que necesitan un docx completo generado de extremo a
    extremo con un driver falso."""
    rows_actual = [
        {
            "estacion_id": "E3",
            "estacion_nombre": "AOI de gama básica",
            "maquina_id": "M3",
            "capacidad_ud_h": 400.0,
            "oee_actual": 0.75,
            "registros": [
                {"fecha": "2026-06-05", "unidades": 4800, "paradas_min": 0.0, "defectos": 50},
                {"fecha": "2026-06-06", "unidades": 4700, "paradas_min": 20.0, "defectos": 60},
            ],
        },
        {
            "estacion_id": "E4",
            "estacion_nombre": "Ensamblaje de carcasa",
            "maquina_id": "M4",
            "capacidad_ud_h": 500.0,
            "oee_actual": 0.88,
            "registros": [
                {"fecha": "2026-06-05", "unidades": 7000, "paradas_min": 10.0, "defectos": 30},
                {"fecha": "2026-06-06", "unidades": 7100, "paradas_min": 15.0, "defectos": 25},
            ],
        },
    ]
    rows_anterior = [
        {
            "estacion_id": "E3",
            "estacion_nombre": "AOI de gama básica",
            "maquina_id": "M3",
            "capacidad_ud_h": 400.0,
            "oee_actual": 0.75,
            "registros": [
                {"fecha": "2026-05-05", "unidades": 4600, "paradas_min": 10.0, "defectos": 55},
            ],
        },
        {
            "estacion_id": "E4",
            "estacion_nombre": "Ensamblaje de carcasa",
            "maquina_id": "M4",
            "capacidad_ud_h": 500.0,
            "oee_actual": 0.88,
            "registros": [
                {"fecha": "2026-05-05", "unidades": 6900, "paradas_min": 12.0, "defectos": 28},
            ],
        },
    ]
    return rows_actual, rows_anterior


def _construir_driver_informe_fixture() -> _FakeDriver:
    """Driver falso con la cola de respuestas que espera _generar_informe:
    resolver línea (periodo actual) + rows actual + rows anterior (línea
    ya resuelta) + candidatos_sustitucion (máquina + candidatas)."""
    rows_actual, rows_anterior = _construir_datos_informe_fixture()
    return _FakeDriver(
        [
            _FakeResult(single_valor={"id": "L1"}),
            _FakeResult(data_valor=rows_actual),
            _FakeResult(data_valor=rows_anterior),
            _FakeResult(single_valor={"maquina": MAQUINA_M3}),
            _FakeResult(data_valor=[{"candidata": CANDIDATA_C3}]),
        ]
    )


def _extraer_texto_completo(doc: Document) -> str:
    """Todo el texto legible del docx: párrafos del cuerpo (doc.paragraphs
    no incluye texto de dentro de tablas), celdas de las 3 tablas
    (portada, tarjetas, KPIs), y el pie de página. No inspecciona las
    gráficas (son imágenes, no texto)."""
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            partes.extend(celda.text for celda in fila.cells)
    for section in doc.sections:
        partes.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(partes)


def test_generar_informe_produce_docx_valido(tmp_path):
    driver = _construir_driver_informe_fixture()

    resultado = server._generar_informe(driver, "mensual", "2026-07-04", output_dir=tmp_path)

    ruta = Path(resultado.ruta)
    assert ruta.exists()

    doc = Document(str(ruta))
    encabezados = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Resumen ejecutivo" in encabezados
    assert "KPIs por estación" in encabezados
    assert "Cuello de botella" in encabezados
    assert "Recomendaciones" in encabezados

    # 3 tablas: portada (1x1), tarjetas de resumen (1x3), KPIs (5 columnas)
    assert len(doc.tables) == 3
    tabla_portada = doc.tables[0]
    tabla_tarjetas = doc.tables[1]
    tabla = next(t for t in doc.tables if len(t.columns) == 5)
    assert len(tabla_portada.rows) == 1 and len(tabla_portada.columns) == 1
    assert len(tabla.rows) == 3  # cabecera + 2 estaciones
    assert len(tabla_tarjetas.columns) == 3

    todos_los_parrafos = [p.text for p in doc.paragraphs]
    idx_cuello = todos_los_parrafos.index("Cuello de botella")
    diagnostico_texto = todos_los_parrafos[idx_cuello + 1]
    idx_recomendaciones = todos_los_parrafos.index("Recomendaciones")
    recomendaciones_texto = todos_los_parrafos[idx_recomendaciones + 1 :]

    # el diagnóstico de "Cuello de botella" no se repite en "Recomendaciones"
    assert diagnostico_texto not in recomendaciones_texto

    texto_completo = _extraer_texto_completo(doc)

    # nombre legible, nunca el id interno de la máquina
    assert "M3" not in texto_completo
    assert "M4" not in texto_completo

    # pie de página con número de página automático (no texto fijo)
    pie = doc.sections[0].footer.paragraphs[0]
    assert "NEXATRON Electronics" in pie.text
    assert "Línea L1" in pie.text
    campos_pagina = pie._p.findall(".//" + qn("w:fldChar"))
    assert len(campos_pagina) == 2  # begin + end del campo PAGE

    # números grandes con separador de miles en formato español
    assert "9.500" in texto_completo  # unidades totales E3 (4800+4700)
    assert "14.100" in texto_completo  # unidades totales E4 (7000+7100)
    assert "23.600" in texto_completo  # unidades totales de la línea


def test_generar_informe_contenido_sin_duplicacion_ids_ni_numeros_mal_formateados(tmp_path):
    """Regresión de contenido, no solo de estructura: los 3 bugs reales
    encontrados en este proyecto (texto duplicado entre diagnóstico y
    recomendaciones, ids internos filtrados al texto de usuario, números
    sin separador de miles) se detectaron abriendo el docx a mano — este
    test los atrapa automáticamente."""
    driver = _construir_driver_informe_fixture()

    resultado = server._generar_informe(driver, "mensual", "2026-07-04", output_dir=tmp_path)
    doc = Document(resultado.ruta)

    # 1. el diagnóstico de "Cuello de botella" no se repite (ni total ni
    # parcialmente) en ninguna frase de "Recomendaciones"
    parrafos = [p.text for p in doc.paragraphs]
    diagnostico = parrafos[parrafos.index("Cuello de botella") + 1]
    recomendaciones = parrafos[parrafos.index("Recomendaciones") + 1 :]
    for r in recomendaciones:
        assert r != diagnostico
        assert diagnostico not in r
        assert r not in diagnostico

    texto_completo = _extraer_texto_completo(doc)

    # 2. ningún id interno (letra + número: M3, E2, C1...) en el texto de
    # cara al usuario. linea_id ("L1") es la única excepción legítima: se
    # muestra a propósito en la portada/pie de página, no es un id de
    # máquina/estación/candidata filtrado por error. Confirmado leyendo
    # _construir_documento: la tabla de KPIs usa estacion_nombre, nunca un
    # id, así que no hace falta excluir ninguna tabla de este check.
    ids_encontrados = [m for m in re.findall(r"\b[A-Z]\d+\b", texto_completo) if m != "L1"]
    assert ids_encontrados == []

    # 3. ningún número de 4+ cifras sin separador de miles en formato
    # español. Se excluyen las fechas ISO (p.ej. "2026-07-04" en la
    # portada) porque un año de 4 dígitos no lleva separador y daría un
    # falso positivo.
    texto_sin_fechas = re.sub(r"\d{4}-\d{2}-\d{2}", "", texto_completo)
    assert re.findall(r"\b\d{4,}\b", texto_sin_fechas) == []


def test_ejecutar_generar_informe_periodo_invalido_devuelve_error_estructurado():
    resultado = server._ejecutar_generar_informe(_FakeDriver([]), "semanal", "2026-07-04")

    assert "periodo debe ser uno de" in resultado["error"]
