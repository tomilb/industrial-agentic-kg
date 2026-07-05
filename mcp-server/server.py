"""Servidor MCP (stdio) con la tool consultar_grafo sobre el grafo de
planta en Neo4j. Contrato de la tool documentado en ARCHITECTURE.md.
"""

from __future__ import annotations

import os
import tempfile
from collections.abc import Callable
from datetime import date, timedelta
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")  # sin display: imprescindible en un servidor MCP headless
import matplotlib.pyplot as plt
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP
from neo4j import Driver, GraphDatabase
from pydantic import BaseModel, Field, ValidationError

# --- Conexión a Neo4j -------------------------------------------------------


def get_driver() -> Driver:
    load_dotenv()
    uri = os.environ["NEO4J_URI"]
    user = os.environ["NEO4J_USER"]
    password = os.environ["NEO4J_PASSWORD"]
    return GraphDatabase.driver(uri, auth=(user, password))


# --- Modelos Pydantic (contratos de ARCHITECTURE.md) ------------------------


class MaquinaSpecs(BaseModel):
    id: str
    nombre: str
    capacidad_ud_h: float
    coste_adquisicion: float
    coste_mantenimiento_anual: float
    consumo_kwh: float
    vida_util_anos: int
    oee_actual: float | None = None
    mejora_capacidad_pct: float | None = None
    anos_en_servicio: int | None = None  # solo en Maquina, no en MaquinaCandidata


class SpecsMaquinaResultado(BaseModel):
    maquina: MaquinaSpecs


class EstacionTopologia(BaseModel):
    id: str
    nombre: str
    maquina_id: str | None = None


class TopologiaLineaResultado(BaseModel):
    linea_id: str
    estaciones: list[EstacionTopologia]


class CandidatosSustitucionResultado(BaseModel):
    maquina_actual: MaquinaSpecs
    candidatas: list[MaquinaSpecs]


class ErrorMCP(BaseModel):
    error: str
    detalle: str | None = None


class CalcularFinancieroInput(BaseModel):
    coste_inversion: float = Field(gt=0)
    capacidad_actual_ud_h: float = Field(ge=0)
    capacidad_nueva_ud_h: float = Field(ge=0)
    margen_por_unidad: float = Field(gt=0)
    horas_operacion_dia: float = Field(gt=0)
    horizonte_anos: int = Field(gt=0)


class ResultadoFinanciero(BaseModel):
    payback_meses: float | None
    roi_pct: float
    van: float


class EstacionCuelloBotella(BaseModel):
    estacion_id: str
    estacion_nombre: str
    maquina_id: str
    capacidad_ud_h: float
    oee_actual: float
    capacidad_efectiva: float
    oee: float | None
    utilizacion_pct: float | None
    fuente_oee: str
    es_restriccion: bool


class DetectarCuelloBotellaResultado(BaseModel):
    linea_id: str
    estaciones: list[EstacionCuelloBotella]


class GenerarInformeResultado(BaseModel):
    ruta: str


# --- Funciones de consulta (puras: reciben el driver, testables con uno falso) --


def _consultar_specs_maquina(driver: Driver, params: dict[str, Any]) -> SpecsMaquinaResultado:
    maquina_id = params.get("maquina_id")
    if not maquina_id:
        raise ValueError("Falta el parámetro 'maquina_id'")

    with driver.session() as session:
        record = session.run(
            """
            MATCH (m {id: $maquina_id})
            WHERE m:Maquina OR m:MaquinaCandidata
            RETURN m{.*} AS maquina
            """,
            maquina_id=maquina_id,
        ).single()

    if record is None:
        raise ValueError(f"No existe ninguna máquina con id {maquina_id!r}")

    return SpecsMaquinaResultado(maquina=MaquinaSpecs(**record["maquina"]))


def _ordenar_estaciones_por_cadena(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Reconstruye el orden de las estaciones a partir de PRECEDE_A: empieza
    por la fila sin predecesora y sigue la cadena predecesora -> siguiente."""
    por_predecesora = {row["predecesora_id"]: row for row in rows}
    ordenadas = []
    actual = por_predecesora.get(None)
    while actual is not None:
        ordenadas.append(
            {
                "id": actual["id"],
                "nombre": actual["nombre"],
                "maquina_id": actual["maquina_id"],
            }
        )
        actual = por_predecesora.get(actual["id"])
    return ordenadas


def _resolver_linea_id(session, linea_id: str | None) -> str:
    """Si no viene linea_id, usa la única Linea del grafo (MVP de una sola línea)."""
    if linea_id is not None:
        return linea_id
    registro = session.run("MATCH (l:Linea) RETURN l.id AS id").single()
    if registro is None:
        raise ValueError("No hay ninguna Linea cargada en el grafo")
    return registro["id"]


def _consultar_topologia_linea(driver: Driver, params: dict[str, Any]) -> TopologiaLineaResultado:
    linea_id = params.get("linea_id")

    with driver.session() as session:
        linea_id = _resolver_linea_id(session, linea_id)

        rows = session.run(
            """
            MATCH (l:Linea {id: $linea_id})-[:TIENE_ESTACION]->(e:Estacion)
            OPTIONAL MATCH (e)-[:OPERA]->(m:Maquina)
            OPTIONAL MATCH (predecesora:Estacion)-[:PRECEDE_A]->(e)
            RETURN e.id AS id, e.nombre AS nombre, m.id AS maquina_id,
                   predecesora.id AS predecesora_id
            """,
            linea_id=linea_id,
        ).data()

    if not rows:
        raise ValueError(f"No existe ninguna Linea con id {linea_id!r}")

    return TopologiaLineaResultado(
        linea_id=linea_id,
        estaciones=[EstacionTopologia(**e) for e in _ordenar_estaciones_por_cadena(rows)],
    )


def _consultar_candidatos_sustitucion(
    driver: Driver, params: dict[str, Any]
) -> CandidatosSustitucionResultado:
    maquina_id = params.get("maquina_id")
    if not maquina_id:
        raise ValueError("Falta el parámetro 'maquina_id'")

    with driver.session() as session:
        record = session.run(
            "MATCH (m:Maquina {id: $maquina_id}) RETURN m{.*} AS maquina",
            maquina_id=maquina_id,
        ).single()
        if record is None:
            raise ValueError(f"No existe ninguna Maquina con id {maquina_id!r}")

        candidatas_rows = session.run(
            """
            MATCH (m:Maquina {id: $maquina_id})-[:PUEDE_SUSTITUIRSE_POR]->(c:MaquinaCandidata)
            RETURN c{.*} AS candidata
            """,
            maquina_id=maquina_id,
        ).data()

    return CandidatosSustitucionResultado(
        maquina_actual=MaquinaSpecs(**record["maquina"]),
        candidatas=[MaquinaSpecs(**row["candidata"]) for row in candidatas_rows],
    )


# --- Dispatcher (separado del decorador MCP para poder testearlo directo) --

PREGUNTAS_VALIDAS = {"specs_maquina", "topologia_linea", "candidatos_sustitucion"}

_CONSULTAS: dict[str, Callable[[Driver, dict[str, Any]], BaseModel]] = {
    "specs_maquina": _consultar_specs_maquina,
    "topologia_linea": _consultar_topologia_linea,
    "candidatos_sustitucion": _consultar_candidatos_sustitucion,
}


def _ejecutar_consulta(driver: Driver, pregunta_estructurada: str, params: dict[str, Any]) -> dict:
    try:
        if pregunta_estructurada not in PREGUNTAS_VALIDAS:
            return ErrorMCP(
                error="pregunta_estructurada no reconocida",
                detalle=f"Recibido {pregunta_estructurada!r}; válidas: {sorted(PREGUNTAS_VALIDAS)}",
            ).model_dump()

        resultado = _CONSULTAS[pregunta_estructurada](driver, params)
        return resultado.model_dump()
    except ValueError as e:
        return ErrorMCP(error=str(e)).model_dump()
    except Exception as e:  # neo4j.exceptions.* u otros fallos inesperados
        return ErrorMCP(error="Error al consultar el grafo", detalle=str(e)).model_dump()


# --- Cálculo financiero (puro: sin Neo4j) -----------------------------------

TASA_DESCUENTO_VAN = 0.10  # ver ARCHITECTURE.md


def _calcular_financiero(entrada: CalcularFinancieroInput) -> ResultadoFinanciero:
    ganancia_diaria = (
        (entrada.capacidad_nueva_ud_h - entrada.capacidad_actual_ud_h)
        * entrada.margen_por_unidad
        * entrada.horas_operacion_dia
    )

    # ganancia_diaria <= 0 (la nueva capacidad no mejora o empeora) no es un
    # error: es una inversión que nunca se recupera. payback_meses queda en
    # None en vez de dividir por cero o devolver un valor negativo sin sentido.
    payback_meses = (
        entrada.coste_inversion / (ganancia_diaria * 30) if ganancia_diaria > 0 else None
    )

    roi_pct = (
        ((ganancia_diaria * 365 * entrada.horizonte_anos) - entrada.coste_inversion)
        / entrada.coste_inversion
        * 100
    )

    flujo_anual = ganancia_diaria * 365
    van = -entrada.coste_inversion + sum(
        flujo_anual / (1 + TASA_DESCUENTO_VAN) ** anio
        for anio in range(1, entrada.horizonte_anos + 1)
    )

    return ResultadoFinanciero(payback_meses=payback_meses, roi_pct=roi_pct, van=van)


def _ejecutar_calculo_financiero(kwargs: dict[str, Any]) -> dict:
    try:
        entrada = CalcularFinancieroInput(**kwargs)
        return _calcular_financiero(entrada).model_dump()
    except ValidationError as e:
        return ErrorMCP(error="Parámetros de entrada inválidos", detalle=str(e)).model_dump()
    except Exception as e:
        return ErrorMCP(
            error="Error al calcular la viabilidad financiera", detalle=str(e)
        ).model_dump()


# --- Detección de cuello de botella ------------------------------------------

# Fijo, ver docs/ESCENARIO.md — no es parámetro de la tool porque el
# histórico ya se generó con este valor fijo.
HORAS_OPERACION_DIA = 16


def _validar_periodo(desde: str, hasta: str) -> None:
    try:
        desde_d = date.fromisoformat(desde)
        hasta_d = date.fromisoformat(hasta)
    except ValueError as e:
        raise ValueError(f"Fecha de periodo inválida (usar YYYY-MM-DD): {e}") from e
    if desde_d > hasta_d:
        raise ValueError(f"'desde' ({desde}) es posterior a 'hasta' ({hasta})")


def _calcular_oee_y_utilizacion(
    registros: list[dict[str, Any]], capacidad_ud_h: float
) -> tuple[float | None, float | None]:
    """OEE y utilización empíricos a partir de los RegistroProduccion reales
    del periodo. None si no hay registros (no es un error, falta de datos)."""
    if not registros:
        return None, None

    tiempo_disponible_min = HORAS_OPERACION_DIA * 60
    oees = [r["unidades"] / (capacidad_ud_h * HORAS_OPERACION_DIA) for r in registros]
    utilizaciones = [
        (tiempo_disponible_min - r["paradas_min"]) / tiempo_disponible_min * 100 for r in registros
    ]
    return sum(oees) / len(oees), sum(utilizaciones) / len(utilizaciones)


def _resolver_oee_para_capacidad(
    oee_empirico: float | None, oee_actual: float
) -> tuple[float, str]:
    """El oee empírico del periodo manda; si no hay registros (None), cae a
    oee_actual (ficha estática) solo para esa estación."""
    if oee_empirico is not None:
        return oee_empirico, "empirico"
    return oee_actual, "estatico_sin_datos"


def _consultar_rows_cuello_botella(
    driver: Driver, linea_id: str | None, desde: str, hasta: str
) -> tuple[str, list[dict[str, Any]]]:
    """Datos crudos por estación (specs + registros del periodo). Se separa
    de _detectar_cuello_botella para que generar_informe pueda reutilizarla
    sin repetir el Cypher (registros incluye fecha/defectos para las
    gráficas y KPIs del informe, no solo unidades/paradas_min)."""
    with driver.session() as session:
        linea_id = _resolver_linea_id(session, linea_id)

        rows = session.run(
            """
            MATCH (l:Linea {id: $linea_id})-[:TIENE_ESTACION]->(e:Estacion)-[:OPERA]->(m:Maquina)
            OPTIONAL MATCH (m)-[:REGISTRA]->(r:RegistroProduccion)
            WHERE r.fecha >= date($desde) AND r.fecha <= date($hasta)
            RETURN e.id AS estacion_id, e.nombre AS estacion_nombre, m.id AS maquina_id,
                   m.capacidad_ud_h AS capacidad_ud_h, m.oee_actual AS oee_actual,
                   collect(r{.unidades, .paradas_min, .defectos,
                             fecha: toString(r.fecha)}) AS registros
            """,
            linea_id=linea_id,
            desde=desde,
            hasta=hasta,
        ).data()

    if not rows:
        raise ValueError(f"No existe ninguna Linea con id {linea_id!r}")

    return linea_id, rows


def _calcular_estaciones_cuello_botella(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    estaciones = []
    for row in rows:
        oee, utilizacion_pct = _calcular_oee_y_utilizacion(row["registros"], row["capacidad_ud_h"])
        # capacidad_efectiva usa el oee EMPÍRICO del periodo (rendimiento real
        # de la estación) cuando hay datos; si no hay registros, cae a
        # oee_actual (ficha estática) solo para esa estación. Nunca decide por
        # capacidad_ud_h nominal sola.
        oee_para_capacidad, fuente_oee = _resolver_oee_para_capacidad(oee, row["oee_actual"])
        capacidad_efectiva = row["capacidad_ud_h"] * oee_para_capacidad
        estaciones.append(
            {
                "estacion_id": row["estacion_id"],
                "estacion_nombre": row["estacion_nombre"],
                "maquina_id": row["maquina_id"],
                "capacidad_ud_h": row["capacidad_ud_h"],
                "oee_actual": row["oee_actual"],
                "capacidad_efectiva": capacidad_efectiva,
                "oee": oee,
                "utilizacion_pct": utilizacion_pct,
                "fuente_oee": fuente_oee,
                "es_restriccion": False,
            }
        )

    capacidad_minima = min(e["capacidad_efectiva"] for e in estaciones)
    for e in estaciones:
        e["es_restriccion"] = e["capacidad_efectiva"] == capacidad_minima

    return estaciones


def _detectar_cuello_botella(
    driver: Driver, linea_id: str | None, desde: str, hasta: str
) -> DetectarCuelloBotellaResultado:
    _validar_periodo(desde, hasta)
    linea_id, rows = _consultar_rows_cuello_botella(driver, linea_id, desde, hasta)
    return DetectarCuelloBotellaResultado(
        linea_id=linea_id,
        estaciones=[EstacionCuelloBotella(**e) for e in _calcular_estaciones_cuello_botella(rows)],
    )


def _ejecutar_deteccion_cuello_botella(
    driver: Driver, linea_id: str | None, desde: str, hasta: str
) -> dict:
    try:
        return _detectar_cuello_botella(driver, linea_id, desde, hasta).model_dump()
    except ValueError as e:
        return ErrorMCP(error=str(e)).model_dump()
    except Exception as e:
        return ErrorMCP(error="Error al detectar el cuello de botella", detalle=str(e)).model_dump()


# --- Generación de informe ---------------------------------------------------

MARGEN_POR_UNIDAD_EUR = 12.0  # ver docs/ESCENARIO.md, fijo para el demo
HORIZONTE_ANOS_RECOMENDACION = 3  # ver docs/EVAL_QUESTIONS.md ("ROI a 3 años")
PERIODO_DIAS = {"mensual": 30, "trimestral": 91}
RUTA_LOGO_PNG = (
    Path(__file__).parent.parent / "skills" / "informe-corporativo" / "assets" / "logo.png"
)
RUTA_INFORMES_DIR = Path(__file__).parent.parent / "reports" / "output"


def _calcular_rango_periodo(periodo: str, fecha_referencia: date) -> tuple[date, date]:
    if periodo not in PERIODO_DIAS:
        raise ValueError(f"periodo debe ser uno de {sorted(PERIODO_DIAS)}, recibido {periodo!r}")
    hasta = fecha_referencia
    desde = hasta - timedelta(days=PERIODO_DIAS[periodo] - 1)
    return desde, hasta


def _calcular_rango_periodo_anterior(desde: date, hasta: date) -> tuple[date, date]:
    """Ventana contigua anterior, misma duración, sin solape."""
    duracion_dias = (hasta - desde).days + 1
    hasta_anterior = desde - timedelta(days=1)
    desde_anterior = hasta_anterior - timedelta(days=duracion_dias - 1)
    return desde_anterior, hasta_anterior


def _serie_produccion_diaria(rows: list[dict[str, Any]]) -> list[tuple[str, int]]:
    """Unidades totales de la línea por fecha, sumando todas las estaciones."""
    totales: dict[str, int] = {}
    for row in rows:
        for r in row["registros"]:
            totales[r["fecha"]] = totales.get(r["fecha"], 0) + r["unidades"]
    return sorted(totales.items())


def _calcular_kpis_estacion(
    rows: list[dict[str, Any]], estaciones: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    por_id = {e["estacion_id"]: e for e in estaciones}
    kpis = []
    for row in rows:
        estacion = por_id[row["estacion_id"]]
        kpis.append(
            {
                "estacion_nombre": row["estacion_nombre"],
                "utilizacion_pct": estacion["utilizacion_pct"],
                "oee": estacion["oee"],
                "unidades_totales": sum(r["unidades"] for r in row["registros"]),
                "defectos_totales": sum(r["defectos"] for r in row["registros"]),
            }
        )
    return kpis


COLOR_NAVY = "#0B2545"
COLOR_TEAL = "#12C2A9"


def _aplicar_estilo_grafico(ax) -> None:
    """Estilo visual consistente en las 3 gráficas: sin gridlines de fondo,
    sin marco superior/derecho."""
    ax.grid(False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)


def _grafico_produccion_diaria(serie: list[tuple[str, int]], ruta: Path) -> None:
    fechas, unidades = zip(*serie, strict=True)
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.plot(fechas, unidades, marker="o", markersize=3, color=COLOR_NAVY)
    ax.set_title("Producción diaria de la línea")
    paso = max(1, len(fechas) // 8)
    ax.set_xticks(list(fechas[::paso]))
    ax.tick_params(axis="x", rotation=45)
    _aplicar_estilo_grafico(ax)
    fig.tight_layout()
    fig.savefig(ruta, dpi=150)
    plt.close(fig)


def _grafico_oee_por_estacion(estaciones: list[dict[str, Any]], ruta: Path) -> None:
    nombres = [e["estacion_nombre"] for e in estaciones]
    oees = [e["oee"] if e["oee"] is not None else e["oee_actual"] for e in estaciones]
    fig, ax = plt.subplots(figsize=(6, 3))
    barras = ax.bar(nombres, oees, color=COLOR_TEAL)
    ax.bar_label(barras, labels=[_formato_numero_es(v, 2) for v in oees], padding=3)
    ax.set_title("OEE por estación (periodo)")
    ax.tick_params(axis="x", rotation=30)
    _aplicar_estilo_grafico(ax)
    fig.tight_layout()
    fig.savefig(ruta, dpi=150)
    plt.close(fig)


def _grafico_evolucion_restriccion(
    estaciones_actual: list[dict[str, Any]],
    estaciones_anterior: list[dict[str, Any]],
    ruta: Path,
) -> None:
    nombres = [e["estacion_nombre"] for e in estaciones_actual]
    por_id_anterior = {e["estacion_id"]: e for e in estaciones_anterior}
    actuales = [e["capacidad_efectiva"] for e in estaciones_actual]
    anteriores = [
        por_id_anterior[e["estacion_id"]]["capacidad_efectiva"] for e in estaciones_actual
    ]
    x = list(range(len(nombres)))
    fig, ax = plt.subplots(figsize=(6, 3))
    barras_anteriores = ax.bar(
        [i - 0.2 for i in x], anteriores, width=0.4, label="Periodo anterior", color=COLOR_TEAL
    )
    barras_actuales = ax.bar(
        [i + 0.2 for i in x], actuales, width=0.4, label="Periodo actual", color=COLOR_NAVY
    )
    ax.bar_label(barras_anteriores, labels=[_formato_numero_es(v) for v in anteriores], padding=3)
    ax.bar_label(barras_actuales, labels=[_formato_numero_es(v) for v in actuales], padding=3)
    ax.set_xticks(x)
    ax.set_xticklabels(nombres, rotation=30)
    ax.set_title("Capacidad efectiva: periodo actual vs anterior")
    ax.legend()
    _aplicar_estilo_grafico(ax)
    fig.tight_layout()
    fig.savefig(ruta, dpi=150)
    plt.close(fig)


def _formato_numero_es(numero: float, decimales: int = 0) -> str:
    """Formatea un número en estilo español: punto como separador de miles,
    coma como separador decimal (p.ej. 2234516 -> "2.234.516", 0.1 -> "0,1")."""
    signo = "-" if numero < 0 else ""
    entero, _, decimal = f"{abs(numero):.{decimales}f}".partition(".")
    grupos: list[str] = []
    while len(entero) > 3:
        grupos.insert(0, entero[-3:])
        entero = entero[:-3]
    grupos.insert(0, entero)
    parte_entera = ".".join(grupos)
    return f"{signo}{parte_entera},{decimal}" if decimal else f"{signo}{parte_entera}"


def _generar_diagnostico_y_recomendaciones(
    driver: Driver,
    estaciones_actual: list[dict[str, Any]],
    estaciones_anterior: list[dict[str, Any]],
) -> tuple[str, list[str]]:
    """Diagnóstico de la restricción (para la sección "Cuello de botella") y
    recomendaciones — solo conclusiones NUEVAS, no repiten el diagnóstico:
    1) tendencia vs periodo anterior, 2) financiero (si hay candidata) u OEE
    vs ficha (si no). Todo derivado de datos reales, nunca texto de
    plantilla fijo — ver docs/DECISIONS.md para el razonamiento completo."""
    ordenadas = sorted(estaciones_actual, key=lambda e: e["capacidad_efectiva"])
    restriccion = ordenadas[0]
    segunda = ordenadas[1] if len(ordenadas) > 1 else None

    if segunda is not None:
        gap_pct = (
            (segunda["capacidad_efectiva"] - restriccion["capacidad_efectiva"])
            / restriccion["capacidad_efectiva"]
            * 100
        )
        diagnostico = (
            f"La estación {restriccion['estacion_nombre']} es la restricción de la línea, con una "
            f"capacidad efectiva de {_formato_numero_es(restriccion['capacidad_efectiva'])} ud/h "
            f"({restriccion['fuente_oee']}), un {_formato_numero_es(gap_pct)}% por debajo de la "
            f"siguiente estación más lenta ({segunda['estacion_nombre']}, "
            f"{_formato_numero_es(segunda['capacidad_efectiva'])} ud/h)."
        )
    else:
        diagnostico = (
            f"La estación {restriccion['estacion_nombre']} es la restricción de la línea, con una "
            f"capacidad efectiva de {_formato_numero_es(restriccion['capacidad_efectiva'])} ud/h "
            f"({restriccion['fuente_oee']})."
        )

    recomendaciones: list[str] = []

    restriccion_anterior = (
        min(estaciones_anterior, key=lambda e: e["capacidad_efectiva"])
        if estaciones_anterior
        else None
    )
    if restriccion_anterior is not None:
        if restriccion_anterior["estacion_id"] != restriccion["estacion_id"]:
            recomendaciones.append(
                f"La restricción ha cambiado respecto al periodo anterior: era "
                f"{restriccion_anterior['estacion_nombre']} "
                f"({_formato_numero_es(restriccion_anterior['capacidad_efectiva'])} ud/h) "
                f"y ahora es {restriccion['estacion_nombre']} "
                f"({_formato_numero_es(restriccion['capacidad_efectiva'])} ud/h)."
            )
        else:
            correspondiente_anterior = next(
                e for e in estaciones_anterior if e["estacion_id"] == restriccion["estacion_id"]
            )
            delta_pct = (
                (restriccion["capacidad_efectiva"] - correspondiente_anterior["capacidad_efectiva"])
                / correspondiente_anterior["capacidad_efectiva"]
                * 100
            )
            tendencia = "mejorado" if delta_pct > 0 else "empeorado"
            recomendaciones.append(
                f"{restriccion['estacion_nombre']} sigue siendo la restricción; su capacidad "
                f"efectiva ha {tendencia} un {_formato_numero_es(abs(delta_pct))}% respecto al "
                f"periodo anterior "
                f"({_formato_numero_es(correspondiente_anterior['capacidad_efectiva'])} -> "
                f"{_formato_numero_es(restriccion['capacidad_efectiva'])} ud/h)."
            )

    candidatos = _consultar_candidatos_sustitucion(
        driver, {"maquina_id": restriccion["maquina_id"]}
    )
    if candidatos.candidatas:
        candidata = candidatos.candidatas[0]
        maquina_actual_nombre = candidatos.maquina_actual.nombre
        # Se asume que la candidata rendiría con el mismo ratio de OEE que la
        # máquina actual (MaquinaCandidata no tiene oee propio en el esquema).
        oee_usado = restriccion["capacidad_efectiva"] / restriccion["capacidad_ud_h"]
        capacidad_candidata_estimada = candidata.capacidad_ud_h * oee_usado
        # Sustituir la restricción no puede mejorar la línea más allá de lo
        # que permita la SIGUIENTE estación más lenta.
        capacidad_nueva_linea = (
            min(capacidad_candidata_estimada, segunda["capacidad_efectiva"])
            if segunda is not None
            else capacidad_candidata_estimada
        )
        resultado_financiero = _calcular_financiero(
            CalcularFinancieroInput(
                coste_inversion=candidata.coste_adquisicion,
                capacidad_actual_ud_h=restriccion["capacidad_efectiva"],
                capacidad_nueva_ud_h=capacidad_nueva_linea,
                margen_por_unidad=MARGEN_POR_UNIDAD_EUR,
                horas_operacion_dia=HORAS_OPERACION_DIA,
                horizonte_anos=HORIZONTE_ANOS_RECOMENDACION,
            )
        )
        if resultado_financiero.payback_meses is not None:
            # candidata proviene siempre de un nodo MaquinaCandidata real (ver
            # Cypher de _consultar_candidatos_sustitucion), y data-gen/generate.py
            # rellena mejora_capacidad_pct para las 5 candidatas del dataset sin
            # excepción — solo sería None si se cargara una MaquinaCandidata sin
            # ese dato, cosa que el generador actual nunca hace.
            assert candidata.mejora_capacidad_pct is not None, (
                "mejora_capacidad_pct no debería ser None para una MaquinaCandidata "
                "real del grafo"
            )
            recomendaciones.append(
                f"Sustituir {maquina_actual_nombre} por {candidata.nombre} "
                f"(+{_formato_numero_es(candidata.mejora_capacidad_pct)}% de capacidad nominal) "
                f"llevaría la capacidad efectiva de la línea a "
                f"~{_formato_numero_es(capacidad_nueva_linea)} ud/h, con un payback estimado de "
                f"{_formato_numero_es(resultado_financiero.payback_meses, 1)} meses y un ROI a "
                f"{HORIZONTE_ANOS_RECOMENDACION} años de "
                f"{_formato_numero_es(resultado_financiero.roi_pct)}%."
            )
        else:
            recomendaciones.append(
                f"Sustituir {maquina_actual_nombre} por {candidata.nombre} no mejoraría el "
                f"throughput de la línea con los datos actuales (capacidad efectiva estimada tras "
                f"el cambio: ~{_formato_numero_es(capacidad_nueva_linea)} ud/h, frente a los "
                f"{_formato_numero_es(restriccion['capacidad_efectiva'])} ud/h actuales)."
            )
    elif (
        restriccion["fuente_oee"] == "empirico"
        and restriccion["oee_actual"] - restriccion["oee"] >= 0.03
    ):
        recomendaciones.append(
            f"El rendimiento real de {restriccion['estacion_nombre']} "
            f"(oee={_formato_numero_es(restriccion['oee'], 2)}) está "
            f"{_formato_numero_es((restriccion['oee_actual'] - restriccion['oee']) * 100)} puntos "
            f"por debajo de su ficha técnica "
            f"(oee_actual={_formato_numero_es(restriccion['oee_actual'], 2)}); "
            f"antes de invertir en sustitución conviene revisar "
            f"paradas/mantenimiento de esta estación."
        )

    return diagnostico, recomendaciones


COLOR_TARJETA_FONDO = "EAF0F6"


def _rgb_color(color_hex: str) -> RGBColor:
    color_hex = color_hex.lstrip("#")
    return RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))


def _sombrear_celda(celda, color_hex: str) -> None:
    """Aplica un color de fondo a una celda de tabla (python-docx no lo
    expone en su API de alto nivel; hay que construir el XML a mano)."""
    color_hex = color_hex.lstrip("#")
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _anadir_campo_pagina(paragraph) -> None:
    """Inserta el campo de Word "PAGE" (número de página automático), no un
    texto fijo — python-docx tampoco expone esto en su API de alto nivel."""
    run = paragraph.add_run()
    fld_char_inicio = OxmlElement("w:fldChar")
    fld_char_inicio.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_fin = OxmlElement("w:fldChar")
    fld_char_fin.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_inicio)
    run._r.append(instr_text)
    run._r.append(fld_char_fin)


def _agregar_portada(
    doc: DocxDocument, periodo: str, linea_id: str, desde: date, hasta: date
) -> None:
    """Bloque de color navy (tercio superior de la portada) con el logo y el
    título en blanco superpuestos, en vez del logo suelto + título negro."""
    tabla = doc.add_table(rows=1, cols=1)
    tabla.autofit = False
    celda = tabla.cell(0, 0)
    celda.width = Inches(6.5)
    fila = tabla.rows[0]
    fila.height = Inches(2.3)
    fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    _sombrear_celda(celda, COLOR_NAVY)
    celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    parrafo_logo = celda.paragraphs[0]
    parrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo_logo.add_run().add_picture(str(RUTA_LOGO_PNG), width=Inches(1.8))

    parrafo_titulo = celda.add_paragraph()
    parrafo_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = parrafo_titulo.add_run(f"Informe de producción — {periodo.capitalize()}")
    run_titulo.font.size = Pt(24)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    parrafo_subtitulo = celda.add_paragraph()
    parrafo_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_subtitulo = parrafo_subtitulo.add_run(
        f"Línea {linea_id} · {desde.isoformat()} a {hasta.isoformat()} · "
        f"generado el {date.today().isoformat()}"
    )
    run_subtitulo.font.size = Pt(11)
    run_subtitulo.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _agregar_tarjetas_resumen(
    doc: DocxDocument, unidades_totales: int, oee_medio: float, estacion_restriccion: str
) -> None:
    """3 tarjetas en línea (número grande + etiqueta) antes de la tabla de KPIs."""
    tabla = doc.add_table(rows=1, cols=3)
    datos = [
        (_formato_numero_es(unidades_totales), "Unidades producidas"),
        (_formato_numero_es(oee_medio, 2), "OEE medio de la línea"),
        (estacion_restriccion, "Estación restricción"),
    ]
    for celda, (numero, etiqueta) in zip(tabla.rows[0].cells, datos, strict=True):
        _sombrear_celda(celda, COLOR_TARJETA_FONDO)
        celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        parrafo_numero = celda.paragraphs[0]
        parrafo_numero.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_numero = parrafo_numero.add_run(numero)
        run_numero.font.size = Pt(20)
        run_numero.font.bold = True
        run_numero.font.color.rgb = _rgb_color(COLOR_NAVY)

        parrafo_etiqueta = celda.add_paragraph()
        parrafo_etiqueta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_etiqueta = parrafo_etiqueta.add_run(etiqueta)
        run_etiqueta.font.size = Pt(12)
        run_etiqueta.font.bold = True
        run_etiqueta.font.color.rgb = _rgb_color(COLOR_TEAL)


def _agregar_separador(doc: DocxDocument) -> None:
    """Espacio + línea divisoria sutil entre un bloque visual (p.ej. las
    tarjetas de resumen) y el contenido siguiente, para que se lean como
    bloques separados en vez de continuación de la misma tabla."""
    parrafo = doc.add_paragraph()
    parrafo.paragraph_format.space_before = Pt(6)
    parrafo.paragraph_format.space_after = Pt(12)
    pPr = parrafo._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    borde_inferior = OxmlElement("w:bottom")
    borde_inferior.set(qn("w:val"), "single")
    borde_inferior.set(qn("w:sz"), "6")
    borde_inferior.set(qn("w:space"), "1")
    borde_inferior.set(qn("w:color"), "CCCCCC")
    pBdr.append(borde_inferior)
    pPr.append(pBdr)


def _agregar_pie_pagina(doc: DocxDocument, linea_id: str) -> None:
    parrafo = doc.sections[0].footer.paragraphs[0]
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.add_run(f"NEXATRON Electronics · Línea {linea_id} · Página ")
    _anadir_campo_pagina(parrafo)


def _construir_documento(
    linea_id: str,
    periodo: str,
    desde: date,
    hasta: date,
    estaciones: list[dict[str, Any]],
    estaciones_anterior: list[dict[str, Any]],
    serie_diaria: list[tuple[str, int]],
    kpis: list[dict[str, Any]],
    diagnostico: str,
    recomendaciones: list[str],
    output_dir: Path,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        doc = Document()

        _agregar_portada(doc, periodo, linea_id, desde, hasta)

        restriccion = next(e for e in estaciones if e["es_restriccion"])
        unidades_totales = sum(k["unidades_totales"] for k in kpis)
        oee_medio = sum(
            e["oee"] if e["oee"] is not None else e["oee_actual"] for e in estaciones
        ) / len(estaciones)

        doc.add_heading("Resumen ejecutivo", level=1)
        doc.add_paragraph(
            f"La línea produjo {_formato_numero_es(unidades_totales)} unidades en el periodo, con "
            f"un OEE medio de {_formato_numero_es(oee_medio, 2)}. La restricción es "
            f"{restriccion['estacion_nombre']} "
            f"({_formato_numero_es(restriccion['capacidad_efectiva'])} ud/h efectivos)."
        )

        doc.add_heading("KPIs por estación", level=1)
        _agregar_tarjetas_resumen(doc, unidades_totales, oee_medio, restriccion["estacion_nombre"])
        _agregar_separador(doc)

        tabla = doc.add_table(rows=1, cols=5)
        tabla.style = "Table Grid"
        for celda, texto in zip(
            tabla.rows[0].cells,
            ["Estación", "Utilización %", "OEE", "Unidades", "Defectos"],
            strict=True,
        ):
            _sombrear_celda(celda, COLOR_NAVY)
            run = celda.paragraphs[0].add_run(texto)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for kpi in kpis:
            fila = tabla.add_row().cells
            fila[0].text = kpi["estacion_nombre"]
            fila[1].text = (
                _formato_numero_es(kpi["utilizacion_pct"], 1)
                if kpi["utilizacion_pct"] is not None
                else "N/D"
            )
            fila[2].text = _formato_numero_es(kpi["oee"], 2) if kpi["oee"] is not None else "N/D"
            fila[3].text = _formato_numero_es(kpi["unidades_totales"])
            fila[4].text = _formato_numero_es(kpi["defectos_totales"])

        doc.add_heading("Producción diaria", level=2)
        _grafico_produccion_diaria(serie_diaria, tmp_path / "produccion.png")
        doc.add_picture(str(tmp_path / "produccion.png"), width=Inches(6))

        doc.add_heading("OEE por estación", level=2)
        _grafico_oee_por_estacion(estaciones, tmp_path / "oee.png")
        doc.add_picture(str(tmp_path / "oee.png"), width=Inches(6))

        doc.add_heading("Cuello de botella", level=1)
        doc.add_paragraph(diagnostico)
        _grafico_evolucion_restriccion(estaciones, estaciones_anterior, tmp_path / "evolucion.png")
        doc.add_picture(str(tmp_path / "evolucion.png"), width=Inches(6))

        doc.add_heading("Recomendaciones", level=1)
        for r in recomendaciones:
            doc.add_paragraph(r, style="List Bullet")

        _agregar_pie_pagina(doc, linea_id)

        ruta = output_dir / f"informe_{periodo}_{hasta.isoformat()}.docx"
        doc.save(str(ruta))
        return ruta


def _generar_informe(
    driver: Driver,
    periodo: str,
    fecha_referencia: str,
    output_dir: Path = RUTA_INFORMES_DIR,
) -> GenerarInformeResultado:
    try:
        ref = date.fromisoformat(fecha_referencia)
    except ValueError as e:
        raise ValueError(f"fecha_referencia inválida (usar YYYY-MM-DD): {e}") from e

    desde, hasta = _calcular_rango_periodo(periodo, ref)
    desde_anterior, hasta_anterior = _calcular_rango_periodo_anterior(desde, hasta)

    linea_id, rows_actual = _consultar_rows_cuello_botella(
        driver, None, desde.isoformat(), hasta.isoformat()
    )
    estaciones_actual = _calcular_estaciones_cuello_botella(rows_actual)

    serie_diaria = _serie_produccion_diaria(rows_actual)
    if not serie_diaria:
        raise ValueError(
            f"No hay registros de producción entre {desde.isoformat()} y {hasta.isoformat()}"
        )

    _, rows_anterior = _consultar_rows_cuello_botella(
        driver, linea_id, desde_anterior.isoformat(), hasta_anterior.isoformat()
    )
    estaciones_anterior = _calcular_estaciones_cuello_botella(rows_anterior)

    kpis = _calcular_kpis_estacion(rows_actual, estaciones_actual)
    diagnostico, recomendaciones = _generar_diagnostico_y_recomendaciones(
        driver, estaciones_actual, estaciones_anterior
    )

    ruta = _construir_documento(
        linea_id,
        periodo,
        desde,
        hasta,
        estaciones_actual,
        estaciones_anterior,
        serie_diaria,
        kpis,
        diagnostico,
        recomendaciones,
        output_dir,
    )
    return GenerarInformeResultado(ruta=str(ruta))


def _ejecutar_generar_informe(driver: Driver, periodo: str, fecha_referencia: str) -> dict:
    try:
        return _generar_informe(driver, periodo, fecha_referencia).model_dump()
    except ValueError as e:
        return ErrorMCP(error=str(e)).model_dump()
    except Exception as e:
        return ErrorMCP(error="Error al generar el informe", detalle=str(e)).model_dump()


# --- Servidor MCP ------------------------------------------------------------

mcp = FastMCP("industrial-kg")

_driver: Driver | None = None


def _get_shared_driver() -> Driver:
    """Crea el driver en la primera tool call, no al importar el módulo.

    Así los tests pueden importar server.py e inyectar un driver falso en
    las funciones de consulta sin necesitar NEO4J_* en el entorno.
    """
    global _driver
    if _driver is None:
        _driver = get_driver()
    return _driver


@mcp.tool()
def consultar_grafo(pregunta_estructurada: str, params: dict[str, Any] | None = None) -> dict:
    """Consulta el grafo de planta mediante funciones semánticas predefinidas
    (no Cypher libre). Valores válidos de pregunta_estructurada:
    - "specs_maquina": params={"maquina_id": "M3"} -> specs de una Maquina o MaquinaCandidata.
    - "topologia_linea": params={} o {"linea_id": "L1"} -> estaciones de la línea en orden.
    - "candidatos_sustitucion": params={"maquina_id": "M3"} -> máquina actual + sus candidatas.

    Devuelve el JSON del resultado, o {"error": str, "detalle": str | None} si falla.
    """
    return _ejecutar_consulta(_get_shared_driver(), pregunta_estructurada, params or {})


@mcp.tool()
def calcular_financiero(
    coste_inversion: float,
    capacidad_actual_ud_h: float,
    capacidad_nueva_ud_h: float,
    margen_por_unidad: float,
    horas_operacion_dia: float,
    horizonte_anos: int,
) -> dict:
    """Calcula payback (meses), ROI (%) y VAN (tasa de descuento fija 10%
    anual) para sustituir una máquina. Si capacidad_nueva_ud_h no mejora
    sobre capacidad_actual_ud_h, payback_meses es None (nunca se recupera)
    y roi_pct/van salen negativos o nulos, sin lanzar error — es un
    resultado financiero válido, no un fallo. Errores de parámetros (p.ej.
    margen_por_unidad <= 0) devuelven {"error": str, "detalle": str | None}.

    IMPORTANTE: capacidad_actual_ud_h y capacidad_nueva_ud_h deben ser el
    throughput EFECTIVO DE TODA LA LÍNEA (el que marca la estación más
    lenta, antes y después del cambio propuesto) — NO la capacidad de la
    máquina individual que se sustituye. Sustituir una máquina que no es
    la restricción de la línea no mejora el throughput real aunque su
    capacidad individual suba; usar la capacidad de esa máquina sola
    daría un payback/ROI ficticio. Para obtener la capacidad efectiva de
    la línea antes y después, combinar con detectar_cuello_botella (o con
    consultar_grafo + el cálculo capacidad_ud_h * oee de cada estación).
    """
    return _ejecutar_calculo_financiero(
        dict(
            coste_inversion=coste_inversion,
            capacidad_actual_ud_h=capacidad_actual_ud_h,
            capacidad_nueva_ud_h=capacidad_nueva_ud_h,
            margen_por_unidad=margen_por_unidad,
            horas_operacion_dia=horas_operacion_dia,
            horizonte_anos=horizonte_anos,
        )
    )


@mcp.tool()
def detectar_cuello_botella(
    desde: str,
    hasta: str,
    linea_id: str | None = None,
) -> dict:
    """Calcula, por estación, capacidad_efectiva = capacidad_ud_h × oee del
    periodo, y marca es_restriccion=true para la(s) estación(es) con
    capacidad_efectiva mínima de la línea — SIEMPRE por capacidad_efectiva,
    nunca por capacidad_ud_h nominal sola. El oee usado es el EMPÍRICO,
    calculado de los RegistroProduccion reales entre desde y hasta (ISO
    "YYYY-MM-DD"), porque una estación puede rendir peor en la práctica que
    su ficha técnica. Si una estación no tiene registros en ese rango, cae a
    oee_actual (ficha estática) solo para esa estación; fuente_oee indica
    cuál se usó ("empirico" o "estatico_sin_datos"). oee/utilizacion_pct
    salen None cuando no hay registros, sin que eso sea un error.

    Usar la capacidad_efectiva de la estación restricción, antes y
    después de un cambio propuesto, como capacidad_actual_ud_h /
    capacidad_nueva_ud_h al llamar a calcular_financiero.
    """
    return _ejecutar_deteccion_cuello_botella(_get_shared_driver(), linea_id, desde, hasta)


@mcp.tool()
def generar_informe(periodo: str, fecha_referencia: str | None = None) -> dict:
    """Genera el informe corporativo (docx) del periodo indicado, con
    portada, resumen ejecutivo, tabla de KPIs por estación, 3 gráficas
    (producción diaria, OEE por estación, evolución de la restricción vs
    periodo anterior) y recomendaciones — todas derivadas de datos reales
    de detectar_cuello_botella/consultar_grafo/calcular_financiero, nunca
    texto de plantilla fijo. periodo: "mensual" (30 días) o "trimestral"
    (91 días), ventana que termina en fecha_referencia (ISO "YYYY-MM-DD",
    por defecto hoy). Devuelve {"ruta": str} al docx generado en
    reports/output/, o {"error", "detalle"} si falla."""
    ref = fecha_referencia or date.today().isoformat()
    return _ejecutar_generar_informe(_get_shared_driver(), periodo, ref)


if __name__ == "__main__":
    mcp.run()
